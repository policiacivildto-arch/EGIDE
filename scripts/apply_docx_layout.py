from pathlib import Path
from datetime import datetime
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


DOC_TITLE = "TEMPLATE PLANO - OPERAÇÃO +1 DEPARTAMENTO"
DOC_SUBTITLE = "Documento Base para Planejamento Operacional"
DOC_CLASSIFICATION = "USO INTERNO - POLÍCIA CIVIL"
DOC_VERSION = "v1.1"
ORG_NAME = "POLÍCIA CIVIL DO ESTADO DO CEARÁ"
LOGO_CANDIDATES = [
    Path(r"c:\Users\Gigabyte G5 Ke\egide-app\public\images\logo-pcce-horizontal.png"),
    Path(r"c:\Users\Gigabyte G5 Ke\egide-app\public\images\logo-gov.png"),
]

COVER_MARKER_TEXTS = {
    ORG_NAME,
    DOC_CLASSIFICATION,
    DOC_SUBTITLE,
    DOC_TITLE,
}


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_existing_cover(doc: Document) -> None:
    to_remove = []
    for paragraph in doc.paragraphs[:80]:
        text = (paragraph.text or '').strip()
        if text in COVER_MARKER_TEXTS:
            to_remove.append(paragraph)
            continue
        if text.startswith('Atualizado em:'):
            to_remove.append(paragraph)

    for paragraph in to_remove:
        remove_paragraph(paragraph)


def add_cover_page_at_start(doc: Document) -> None:
    if doc.paragraphs:
        anchor = doc.paragraphs[0]
    else:
        anchor = doc.add_paragraph()

    logo_path = next((p for p in LOGO_CANDIDATES if p.exists()), None)
    if logo_path:
        logo_p = anchor.insert_paragraph_before()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_p.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(3.4))

    org = anchor.insert_paragraph_before(ORG_NAME)
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org.runs[0]
    org_run.bold = True
    org_run.font.size = Pt(12)

    footer_note = anchor.insert_paragraph_before(DOC_CLASSIFICATION)
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = footer_note.runs[0]
    c_run.bold = True
    c_run.font.size = Pt(11)

    info = anchor.insert_paragraph_before(f"Atualizado em: {datetime.now().strftime('%d/%m/%Y')}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.runs[0]
    info_run.font.size = Pt(11)

    blank = anchor.insert_paragraph_before("\n\n")
    blank.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = anchor.insert_paragraph_before(DOC_SUBTITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.runs[0]
    r2.font.size = Pt(14)

    title = anchor.insert_paragraph_before(DOC_TITLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.runs[0]
    r.bold = True
    r.font.size = Pt(22)

    page_break = anchor.insert_paragraph_before()
    page_break.add_run().add_break()


def set_footer_all_sections(doc: Document) -> None:
    footer_date = datetime.now().strftime('%d/%m/%Y')
    footer_text = f"{ORG_NAME} | {DOC_TITLE} | {DOC_CLASSIFICATION} | {DOC_VERSION} | {footer_date}"

    for section in doc.sections:
        footer = section.footer
        for existing in list(footer.paragraphs):
            remove_paragraph(existing)
        p = footer.add_paragraph()
        p.text = footer_text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(9)


def main() -> None:
    target = Path(r"c:\Users\Gigabyte G5 Ke\egide-app\_TEMPLATE PLANO - OPERAÇÃO +1 DEPARTAMENTO.docx")
    if not target.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {target}")

    backup = target.with_name(target.stem + "_backup" + target.suffix)
    shutil.copy2(target, backup)

    output = target.with_name(target.stem + "_COM_CAPA_E_RODAPE" + target.suffix)

    doc = Document(str(target))
    strip_existing_cover(doc)
    add_cover_page_at_start(doc)
    set_footer_all_sections(doc)
    saved_target = True
    try:
        doc.save(str(target))
    except PermissionError:
        saved_target = False

    doc.save(str(output))

    if saved_target:
        print(f"OK: layout aplicado em {target}")
    else:
        print(f"AVISO: não foi possível salvar em {target} (arquivo em uso)")
    print(f"Backup: {backup}")
    print(f"Cópia final: {output}")


if __name__ == "__main__":
    main()
